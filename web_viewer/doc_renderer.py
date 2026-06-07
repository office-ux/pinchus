import os
import copy
from typing import List, Dict, Any
import docx

def find_air_outlets_table(doc):
    """Find the table in the document containing 'UNIT AIR OUTLETS'."""
    for table in doc.tables:
        for row in table.rows:  # Look in all rows for header
            for cell in row.cells:
                if "UNIT AIR OUTLETS" in cell.text.upper():
                    return table
    return None

def build_column_mapping(table) -> Dict[int, str]:
    """
    Build a mapping from column index to database field name
    by reading Group Header (Row 15) and Detail Header (Row 16).
    """
    # 1. First, we need to locate the row containing DWG # or VEL/CFM detail headers
    group_row_idx = -1
    detail_row_idx = -1
    
    for r_idx, row in enumerate(table.rows):
        cells_text = [c.text.strip().upper() for c in row.cells]
        if "DWG #" in cells_text or "VEL" in cells_text:
            detail_row_idx = r_idx
            # The row right above is the Group category header row
            if r_idx > 0:
                group_row_idx = r_idx - 1
            break
            
    if detail_row_idx == -1:
        # Fallback to standard indices (15 and 16) if not matched dynamically
        group_row_idx = 15
        detail_row_idx = 16
        
    mapping = {}
    group_row = table.rows[group_row_idx]
    detail_row = table.rows[detail_row_idx]
    
    for i in range(len(detail_row.cells)):
        g_val = group_row.cells[i].text.strip().upper() if group_row_idx >= 0 else ""
        d_val = detail_row.cells[i].text.strip().upper()
        
        # Clean up line breaks
        g_val = g_val.replace("\n", " ").strip()
        d_val = d_val.replace("\n", " ").strip()
        
        field_name = None
        if "DWG #" in d_val:
            field_name = "DWG #"
        elif d_val == "#":
            field_name = "#"
        elif d_val in ("TYPE", "UNIT"):
            field_name = "TYPE"
        elif d_val == "SIZE":
            field_name = "SIZE"
        elif d_val in ("K-FACTOR", "METHOD"):
            field_name = "K-factor"
        elif "DESIGN" in d_val or "DESIGN" in g_val:
            if d_val == "CFM" or "CFM" in d_val:
                field_name = "DESIGN CFM"
        elif "PRELIMINARY" in g_val:
            if "VEL" in d_val:
                field_name = "PRELIMINARY VEL"
            elif "CFM" in d_val:
                field_name = "PRELIMINARY CFM"
            elif "%" in d_val:
                field_name = "PRELIMINARY %"
        elif "FINAL" in g_val:
            if "VEL" in d_val:
                field_name = "FINAL VEL"
            elif "CFM" in d_val:
                field_name = "FINAL CFM"
            elif "%" in d_val:
                field_name = "FINAL %"
                
        if field_name:
            mapping[i] = field_name
            
    return mapping, detail_row_idx

def extract_template_columns(template_path: str) -> List[str]:
    """
    Open the template, find the outlets table, and return the list of detected column names.
    """
    doc = docx.Document(template_path)
    table = find_air_outlets_table(doc)
    if not table:
        raise ValueError("Could not find 'UNIT AIR OUTLETS' table in template document.")
    
    mapping, _ = build_column_mapping(table)
    # The mapping maps column index to field name. We just want a unique list of field names.
    columns = []
    for idx in sorted(mapping.keys()):
        field_name = mapping[idx]
        if field_name not in columns:
            columns.append(field_name)
    return columns


from docx.document import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

def iter_block_items(parent):
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def extract_table_preview(template_path: str) -> dict:
    """
    Open the template and extract all document content in order, 
    preserving paragraphs and rendering tables. Returns the interactive 
    UNIT AIR OUTLETS table separately with detailed structure.
    """
    doc = docx.Document(template_path)
    target_table = find_air_outlets_table(doc)
    if not target_table:
        raise ValueError("Could not find 'UNIT AIR OUTLETS' table in template document.")

    blocks = []
    
    # Pre-process the target table as before
    col_mapping, detail_row_idx = build_column_mapping(target_table)
    group_row_idx = detail_row_idx - 1 if detail_row_idx > 0 else -1
    header_indices = set()
    if group_row_idx >= 0:
        header_indices.add(group_row_idx)
    header_indices.add(detail_row_idx)
    data_start_idx = detail_row_idx + 1

    target_rows_out = []
    global_seen_cells = {}
    for r_idx, row in enumerate(target_table.rows):
        cells_out = []
        col_cursor = 0
        for cell in row.cells:
            cell_id = id(cell)
            if cell_id in global_seen_cells:
                prev_info = global_seen_cells[cell_id]
                if prev_info['start_row'] == r_idx:
                    prev_info['info']['colspan'] += 1
                else:
                    if prev_info['last_seen_row'] < r_idx:
                        prev_info['info']['rowspan'] += 1
                        prev_info['last_seen_row'] = r_idx
                col_cursor += 1
                continue

            cell_info = {
                'text': cell.text.strip(),
                'colspan': 1,
                'rowspan': 1,
                'col_index': col_cursor,
                'is_header': r_idx in header_indices,
                'is_data': r_idx >= data_start_idx,
                'mapped_field': col_mapping.get(col_cursor),
            }
            global_seen_cells[cell_id] = {
                'info': cell_info,
                'start_row': r_idx,
                'last_seen_row': r_idx
            }
            cells_out.append(cell_info)
            col_cursor += 1

        target_rows_out.append({
            'row_index': r_idx,
            'is_header': r_idx in header_indices,
            'is_data': r_idx >= data_start_idx,
            'cells': cells_out,
        })

    col_mapping_out = {str(k): v for k, v in col_mapping.items()}

    # Extract all blocks in order
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text:
                blocks.append({
                    "type": "paragraph",
                    "text": text,
                    "alignment": block.alignment.name if block.alignment else "LEFT"
                })
        elif isinstance(block, Table):
            if block._element == target_table._element:
                blocks.append({
                    "type": "target_table"
                })
            else:
                # Render standard non-interactive table
                static_rows = []
                global_seen_static = {}
                for r_idx, r in enumerate(block.rows):
                    row_cells = []
                    for c in r.cells:
                        c_id = id(c)
                        if c_id in global_seen_static:
                            prev = global_seen_static[c_id]
                            if prev['start_row'] == r_idx:
                                prev['info']['colspan'] += 1
                            else:
                                if prev['last_seen_row'] < r_idx:
                                    prev['info']['rowspan'] += 1
                                    prev['last_seen_row'] = r_idx
                            continue
                        info = {'text': c.text.strip(), 'colspan': 1, 'rowspan': 1}
                        global_seen_static[c_id] = {'info': info, 'start_row': r_idx, 'last_seen_row': r_idx}
                        row_cells.append(info)
                    
                    if row_cells:
                        static_rows.append(row_cells)
                if static_rows:
                    blocks.append({
                        "type": "static_table",
                        "rows": static_rows
                    })

    return {
        'blocks': blocks,
        'target_table': {
            'rows': target_rows_out,
            'header_row_indices': sorted(list(header_indices)),
            'data_start_row': data_start_idx,
            'col_mapping': col_mapping_out,
            'num_cols': len(target_table.columns) if target_table.columns else 0,
        }
    }

def populate_docx_template(template_path: str, output_path: str, stamps_data: List[Dict[str, Any]], user_mapping: Dict[str, str] = None):
    """
    Load .docx template, find outlets table, populate stamps data, and save.
    """
    doc = docx.Document(template_path)
    table = find_air_outlets_table(doc)
    if not table:
        raise ValueError("Could not find 'UNIT AIR OUTLETS' table in template document.")
        
    mapping, header_idx = build_column_mapping(table)
    
    # Template data row is the row immediately following the detail header
    tpl_row_idx = header_idx + 1
    if tpl_row_idx >= len(table.rows):
        raise ValueError("Template does not contain a data row beneath headers.")
        
    tpl_row = table.rows[tpl_row_idx]
    
    # We will copy the template row to insert new rows with the exact format
    rows_to_insert = len(stamps_data)
    
    inserted_rows = []
    for idx, stamp in enumerate(stamps_data):
        # Build field dict
        field_dict = {f["name"]: f["value"] for f in stamp.get("fields", [])}
        
        # We can reuse the existing template row for the first entry
        if idx == 0:
            row = tpl_row
        else:
            # Duplicate the template row
            new_tr = copy.deepcopy(tpl_row._tr)
            # Insert before the TOTAL row (which is at the bottom, or just append)
            # Find insertion index (we insert after the last populated row)
            table._tbl.append(new_tr)
            row = docx.table._Row(new_tr, table)
            
        # Set text in the cells
        for col_idx, doc_col_name in mapping.items():
            # If user provided a mapping, map the document column name to the database field name
            db_field_name = user_mapping.get(doc_col_name, doc_col_name) if user_mapping else doc_col_name
            val = field_dict.get(db_field_name, "")
            # Preserve original fonts and formatting by setting run text rather than cell.text
            cell = row.cells[col_idx]
            # Clear text but keep paragraphs
            if cell.paragraphs:
                p = cell.paragraphs[0]
                p.text = str(val) if val is not None else ""
            else:
                cell.text = str(val) if val is not None else ""
                
        inserted_rows.append(row)
        
    # Delete any original blank template rows between the first row and the TOTAL row
    # In templates, there are typically multiple blank buffer rows up to row 33. We clean them up
    # but leave the TOTAL rows intact at the bottom.
    total_row_idx = -1
    for r_idx, row in enumerate(table.rows):
        first_cell_text = row.cells[0].text.strip().upper()
        if "TOTAL" in first_cell_text:
            total_row_idx = r_idx
            break
            
    # Remove any extra blank rows between tpl_row_idx + rows_to_insert and the total row
    if total_row_idx != -1:
        # Delete rows in reverse order to avoid index shift issues
        start_del = tpl_row_idx + max(1, rows_to_insert)
        end_del = total_row_idx
        for r_idx in range(end_del - 1, start_del - 1, -1):
            if r_idx < len(table.rows):
                tr = table.rows[r_idx]._tr
                table._tbl.remove(tr)
                
    doc.save(output_path)
